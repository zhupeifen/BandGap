"""
Generate a standalone Word report on the six band-gap datasets.

Summarizes each dataset and the cross-dataset findings (zero-inflation, the
two-stage model, the solid-solution generalization mechanism, feature importance),
with the key figures embedded. Output: reports/six_datasets_report.docx.

    .venv/Scripts/python.exe scripts/make_dataset_report.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
PLOTS = ROOT / "plots"
OUT_DIR = ROOT / "reports"
OUT_DIR.mkdir(exist_ok=True)
OUT = OUT_DIR / "six_datasets_report.docx"


def add_table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        c.paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].paragraphs[0].add_run(str(val))
    return t


def fig(doc, name, caption, width=6.2):
    p = PLOTS / name
    if p.exists():
        doc.add_picture(str(p), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)


doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

title = doc.add_heading("Band-Gap Machine Learning: A Six-Dataset Report", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("University of Missouri — Evan Research Group").italic = True

doc.add_heading("1. Overview", level=1)
doc.add_paragraph(
    "We study six public datasets of inorganic band gaps, spanning three orders of "
    "magnitude in size (1.3k–106k materials), the full range of zero-inflation "
    "(0–96% metals), and several target types (PBE, GLLB-SC, experimental, and a fitted "
    "surrogate). All models use Magpie composition descriptors plus each dataset's native "
    "numeric columns, with gradient-boosted trees (XGBoost).")

add_table(doc,
          ["Dataset", "N", "Target", "Non-zero", "Gap range eV (nz median)", "Notes"],
          [["Expt gap", "6,354", "gap expt (experimental)", "61.3%", "0.02–11.70 (1.83)",
            "composition-only; solid solutions present"],
           ["MP gap", "106,113", "gap pbe (PBE)", "56.5%", "0.02–9.72 (1.95)",
            "Materials Project; composition from structures"],
           ["Wolverton oxides", "4,914", "gap pbe (PBE)", "20.8%", "0.19–6.22 (1.45)",
            "ABO3; structural + energetic features"],
           ["Castelli", "18,928", "gap gllbsc (GLLB-SC)", "3.9%", "0.20–7.00 (1.40)",
            "cubic perovskites; electronic features"],
           ["Double perovskites", "1,306", "gap gllbsc (GLLB-SC)", "100%", "0.11–8.34 (4.07)",
            "A2BB'O6; no zero gaps"],
           ["Tol-screened", "67,916", "Band gap(HSE-mf1)", "100%", "0.88–6.52 (3.36)",
            "solid-solution grid; surrogate target; no zeros"]])

doc.add_heading("2. Key cross-dataset findings", level=1)

doc.add_heading("2.1 Aggregate accuracy hides poor regression on semiconductors", level=2)
doc.add_paragraph(
    "Because band-gap datasets are zero-inflated (most candidates are metals), a single "
    "regressor earns a high aggregate accuracy dominated by trivially-correct metals while "
    "doing poorly on the materials that actually have a gap. Reporting non-zero-material MAE "
    "exposes this. On Castelli, aggregate accuracy is 0.93 while the non-zero MAE is ~1 eV.")
add_table(doc,
          ["Dataset (single-stage baseline)", "Aggregate acc <0.07 eV",
           "Non-zero acc <0.07 eV", "Non-zero MAE (eV)"],
          [["Expt gap", "0.317", "0.220", "0.356"],
           ["MP gap (106k)", "0.313", "0.093", "0.601"],
           ["Wolverton", "0.656", "0.204", "0.435"],
           ["Castelli", "0.925", "0.038", "1.013"]])

doc.add_heading("2.2 A two-stage model, and an honest decomposition", level=2)
doc.add_paragraph(
    "A metal/non-metal classifier gating a regressor trained only on non-zero-gap materials "
    "(log target) lowers non-zero MAE by 23–42% end-to-end. But a leak-free nested "
    "cross-validation that holds features constant shows most of that gain comes from the "
    "richer features, not the two-stage structure. The pure method effect is modest where "
    "rich descriptors exist (Wolverton −4.8%, n.s.; Castelli −8.6%, p=0.007) but substantial "
    "on composition-only data (expt_gap −14.7%, p=1e-5; MP gap −7.8%, all folds).")
add_table(doc,
          ["Dataset", "Single-stage MAE", "Two-stage MAE", "Method effect"],
          [["Expt gap (composition-only)", "0.356", "0.304", "−14.7% (p=1e-5)"],
           ["MP gap (106k)", "0.601", "0.554", "−7.8% (all folds)"],
           ["Wolverton", "0.364", "0.347", "−4.8% (n.s.)"],
           ["Castelli", "0.651", "0.595", "−8.6% (p=0.007)"]])
fig(doc, "feature_importance.png",
    "Figure: SHAP feature importance per dataset. Energetic descriptors dominate Wolverton "
    "(51% of total importance); electronegativity statistics drive the experimental set.")

doc.add_heading("2.3 Random-split optimism is concentrated on solid-solution datasets", level=2)
doc.add_paragraph(
    "Random train/test splits can place near-identical compositions on both sides, "
    "overstating generalization. The size of this effect is dataset-dependent. A simple "
    "redundancy count does NOT predict it (Pearson r=0.54, n.s.) — Castelli has 49% repeated "
    "compositions yet no optimism. The factor that does separate the datasets is the share of "
    "solid-solution (fractional) compositions, which let the model interpolate within a "
    "continuous composition family. Optimism is large only on those sets (Tol 99.8% fractional, "
    "expt_gap 21.5%) and negligible on collections of distinct compounds; the solid-solution "
    "share tracks optimism with Spearman 0.90.")
add_table(doc,
          ["Dataset", "Solid-solution share", "Optimism (by-chemistry / random MAE)"],
          [["Tol-screened", "99.8%", "1.64×"],
           ["Expt gap", "21.5%", "1.55×"],
           ["MP gap (106k)", "~0% (polymorph duplicates)", "1.14×"],
           ["Wolverton", "1.3%", "1.02×"],
           ["Castelli", "0.5%", "1.00×"],
           ["Double perovskites", "0.0%", "1.01×"]])
fig(doc, "generalization_splits.png",
    "Figure: MAE under leakage-controlled splits, relative to a random split, for all six "
    "datasets. Optimism is large on the solid-solution sets, mild on MP gap, negligible elsewhere.")
fig(doc, "redundancy_correlation.png",
    "Figure: Random-split optimism vs. solid-solution share (Spearman 0.90). A simple "
    "redundancy count does not predict optimism (r=0.54, n.s.).")

doc.add_heading("2.4 Data-limited, and the no-zero datasets", level=2)
doc.add_paragraph(
    "Learning curves show Castelli and expt_gap are still improving at full training size — "
    "data-limited, not algorithm-limited. The two no-zero datasets (Double perovskites, Tol) "
    "need no two-stage gating: Tol's surrogate target is nearly deterministic (R²≈0.9998, MAE "
    "0.015 eV) and even survives chemistry-extrapolation at R²≈0.999, while the double-perovskite "
    "set (1,306 distinct compounds) regresses to MAE ≈ 0.25 eV with no split optimism.")
fig(doc, "learning_curve.png",
    "Figure: Learning curves. Neither data-limited dataset has plateaued at full size.")

doc.add_heading("3. Takeaways", level=1)
for t in [
    "Report non-zero-material MAE alongside aggregate accuracy on zero-inflated band-gap data; "
    "aggregate accuracy is dominated by metals and is misleading.",
    "The two-stage classify-then-regress model helps most when only composition features are "
    "available; with rich structural/energetic features, the features carry most of the gain.",
    "Run a grouped (by-chemistry) split as a cheap diagnostic — it matters most for datasets "
    "containing solid solutions (alloys, mixed-site screening grids).",
    "Energetic descriptors (formation/hull energy) are the single most informative feature group "
    "for the oxide gaps; electronegativity statistics drive the experimental set.",
]:
    doc.add_paragraph(t, style="List Bullet")

doc.add_paragraph()
note = doc.add_paragraph()
note.add_run("Generated from the analysis scripts in scripts/ (see reproduce_all.py). "
             "All numbers are 5–10-fold cross-validated unless noted.").italic = True

doc.save(OUT)
print(f"Wrote {OUT}  ({OUT.stat().st_size // 1024} KB)")
